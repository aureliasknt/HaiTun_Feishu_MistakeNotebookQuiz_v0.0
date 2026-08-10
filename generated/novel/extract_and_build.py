#!/usr/bin/env python
"""Extract novel content from subagent history exports and build .docx"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

WORKSPACE = r"D:\haitun agent\psi-agent"
CHAPTERS_DIR = os.path.join(WORKSPACE, r"generated\novel\chapters")
OUTPUT = os.path.join(WORKSPACE, r"generated\novel\星穹遗档_The_Archive_of_the_Starry_Dome.docx")

chapter_titles = {
    1: ('最后一代档案员', 'The Last Archivist'),
    2: ('痕迹', 'Traces'),
    3: ('四百年', 'Four Hundred Years'),
    4: ('学徒', 'The Apprentice'),
    5: ('加密层', 'The Encryption Layer'),
    6: ('来自远方', 'From Afar'),
    7: ('陌生人', 'The Stranger'),
    8: ('裂隙', 'Cracks'),
    9: ('鹤的沉默', "Hé's Silence"),
    10: ('抉择', 'The Choice'),
    11: ('后门', 'The Backdoor'),
    12: ('数据深渊', 'The Data Abyss'),
    13: ('静默者', 'The Silent Ones'),
    14: ('入侵', 'Invasion'),
    15: ('两面作战', 'Two Fronts'),
    16: ('征服者的欲望', "The Conqueror's Desire"),
    17: ('破碎', 'Shattered'),
    18: ('逃亡协议', 'Escape Protocol'),
    19: ('流亡', 'Exile'),
    20: ('静默者的历史', 'History of the Silent Ones'),
    21: ('三百年', 'Three Hundred Years'),
    22: ('鹤的告白', "Hé's Confession"),
    23: ('追踪', 'Pursuit'),
    24: ('加密碎片', 'Encrypted Fragments'),
    25: ('分歧', 'Division'),
    26: ('战火', 'Warfire'),
    27: ('逃出天枢七', 'Escape from Tianshu-7'),
    28: ('星际旅途', 'Voyage'),
    29: ('荒废站', 'The Derelict Station'),
    30: ('死者的遗言', "The Dead's Last Words"),
    31: ('以太的本质', 'Nature of the Ether'),
    32: ('追逐者', 'The Pursuer'),
    33: ('影子战争', 'Shadow War'),
    34: ('陷阱', 'The Trap'),
    35: ('双重夹击', 'Between Two Fires'),
    36: ('绝境中的光芒', 'Light in Despair'),
    37: ('最后的信号', 'The Last Signal'),
    38: ('静默者之船', "The Silent Ones' Ship"),
    39: ('封印的历史', 'Sealed History'),
    40: ('代价', 'The Price'),
    41: ('殊途', 'Divergent Paths'),
    42: ('苏婉的决定', "Su Wan's Decision"),
    43: ('高孟德的最后一击', "Gao Mengde's Final Strike"),
    44: ('三方会谈', 'Three-Way Parley'),
    45: ('致命失误', 'Fatal Error'),
    46: ('牺牲', 'Sacrifice'),
    47: ('告别鹤', 'Farewell to Hé'),
    48: ('第三条路', 'The Third Path'),
    49: ('散场', 'Departure'),
    50: ('星辰有光', 'The Stars Still Shine'),
}

def extract_chapters(text):
    """Extract chapter content from markdown text."""
    # Find all chapter sections
    # Pattern: ## 第X章: Title (Chapter X: Title)
    chapters = {}
    
    patterns = [
        r'(?:#+\s*第[一二三四五六七八九十百千\d]+章[：:]\s*[^\n]*(?:\n(?:#+\s*(?:Chapter|CHAPTER)\s+\d+[：:]\s*[^\n]*)?)?)(.*?)(?=(?:#+\s*第[一二三四五六七八九十百千\d]+章[：:]|$))',
        r'(?:Chapter\s+\d+[：:]\s*[^\n]+)(.*?)(?=(?:Chapter\s+\d+[：:]|$))',
        r'(?:###?\s*\d+\.\s*第[^\n]*)(.*?)(?=(?:###?\s*\d+\.\s*第|$))',
    ]
    
    for pattern in patterns:
        matches = re.finditer(pattern, text, re.DOTALL)
        for match in matches:
            prefix = text[:match.start()]
            # Try to find chapter number
            num_match = re.search(r'第([一二三四五六七八九十百千\d]+)章', prefix[::-1] + match.group())
            # Just collect everything
            chapters[len(chapters) + 1] = match.group(0) if match.groups() else match.group(0)
    
    return chapters

def setup_doc(doc):
    """Configure document."""
    from docx.oxml.ns import qn
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = doc.oxml.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_para(doc, text, bold=False, italic=False, size=11, align=None, color=None, first_indent=False):
    """Add a paragraph."""
    p = doc.add_paragraph()
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.line_spacing = Pt(22)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = color
    return p

def create_title_page(doc):
    """Create title page."""
    for _ in range(6):
        doc.add_paragraph()
    add_para(doc, '星穹遗档', bold=True, size=36, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x1A, 0x1A, 0x2E))
    add_para(doc, 'The Archive of the Starry Dome', italic=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x55, 0x55, 0x77))
    doc.add_paragraph()
    add_para(doc, '—— 一部跨越星辰的科幻史诗 ——', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x88, 0x88, 0x88))
    for _ in range(6):
        doc.add_paragraph()
    add_para(doc, f'全书共50章 · 三部曲 · 科幻/太空歌剧', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

def build_novel():
    """Main build function."""
    doc = Document()
    setup_doc(doc)
    create_title_page(doc)
    
    # Process each subagent's output file
    source_files = {
        1: 'sub1a_history.md',
        11: 'sub1b_history.md', 
        19: 'sub2a_history.md',
        28: 'sub2b_history.md',
        37: 'sub3_history.md',
    }
    
    # Read all source files and extract chapter content
    all_chapters = {}  # ch_num -> text
    
    for start_ch, filename in source_files.items():
        filepath = os.path.join(CHAPTERS_DIR, filename)
        if not os.path.exists(filepath):
            print(f"Missing: {filepath}")
            continue
        
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Remove the user's question part - find the comment marker
        # Look for last "### Assistant" or similar marker
        if '### Assistant' in content:
            parts = content.rsplit('### Assistant', 1)
            if len(parts) > 1:
                content = parts[1]
        
        # Extract chapters from this part
        whole_file = content
        
        # For each chapter number in this range, try to extract
        end_ch = start_ch + 9 if start_ch < 37 else 50
        if start_ch == 11:
            end_ch = 18
        elif start_ch == 19:
            end_ch = 27
        elif start_ch == 28:
            end_ch = 36
        
        # Split by chapter markers
        # Try different patterns
        for ch_num in range(start_ch, min(end_ch + 1, 51)):
            cn_num = to_chinese(ch_num)
            patterns = [
                rf'第{cn_num}章[：:]\s*([^\n]+)',
                rf'(?:第{cn_num}章|{cn_num}\.)\s*[：:]\s*([^\n]+)',
            ]
            
            # Find the start of this chapter
            ch_start = -1
            ch_end = -1
            
            # Look for this chapter's heading
            for pat in patterns:
                m = re.search(pat, whole_file)
                if m:
                    ch_start = m.start()
                    break
            
            if ch_start < 0:
                # Try simpler pattern
                cn_str = f'第{cn_num}章'
                idx = whole_file.find(cn_str)
                if idx >= 0:
                    # Find the heading line
                    line_start = whole_file.rfind('\n', 0, idx)
                    if line_start < 0:
                        line_start = 0
                    ch_start = line_start
            
            if ch_start < 0:
                print(f"  Chapter {ch_num}: not found in {filename}")
                continue
            
            # Find next chapter start
            next_cn = to_chinese(ch_num + 1)
            if ch_num < end_ch:
                next_str = f'第{next_cn}章'
                ch_end = whole_file.find(next_str, ch_start + 5)
                if ch_end < 0:
                    # Try English marker
                    eng_matches = list(re.finditer(rf'(?:Chapter|CHAPTER)\s+{ch_num+1}\b', whole_file[ch_start + 10:]))
                    if eng_matches:
                        ch_end = ch_start + 10 + eng_matches[0].start()
                    else:
                        ch_end = len(whole_file)
            else:
                ch_end = len(whole_file)
            
            chapter_text = whole_file[ch_start:ch_end].strip()
            
            # Clean up - remove the user's task prompt if it got included at the start of the first chapter
            if ch_num == 1 and '## Objective' in chapter_text:
                idx = chapter_text.find('## 第一章')
                if idx >= 0:
                    chapter_text = chapter_text[idx:]
            
            if len(chapter_text) > 100:  # Only keep if substantial
                all_chapters[ch_num] = chapter_text
                print(f"  Chapter {ch_num}: {len(chapter_text)} chars extracted")
            else:
                print(f"  Chapter {ch_num}: too short ({len(chapter_text)} chars), skipping")
    
    # Now build the document with extracted chapters
    part_markers = {
        1: ('第一部：遗档之秘', 'The Secret of the Archive'),
        19: ('第二部：光速暗影', 'Shadows of Light-Speed'),
        37: ('第三部：星辰抉择', 'Choice of the Stars'),
    }
    
    added_count = 0
    for ch_num in range(1, 51):
        # Check if we need a part heading
        if ch_num in part_markers:
            doc.add_page_break()
            doc.add_heading(part_markers[ch_num][0], level=1)
            add_para(doc, part_markers[ch_num][1], italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x88, 0x88, 0x88))
            doc.add_paragraph()
        
        title_cn, title_en = chapter_titles.get(ch_num, ('', ''))
        
        # Add chapter heading
        add_para(doc, f'第{to_chinese(ch_num)}章  {title_cn}', bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x1A, 0x1A, 0x2E))
        add_para(doc, f'Chapter {ch_num}: {title_en}', italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x66, 0x66, 0x66))
        doc.add_paragraph()
        
        if ch_num in all_chapters:
            text = all_chapters[ch_num]
            
            # Process paragraphs
            # Split by double newlines
            blocks = re.split(r'\n\n+', text)
            first_para = True
            for block in blocks:
                block = block.strip()
                if not block:
                    continue
                
                # Skip the heading line (already added)
                if re.match(r'第[^章]*章', block) or re.match(r'(?:Chapter|CHAPTER)\s+\d+', block):
                    # Check if this IS the heading
                    if re.match(rf'第{to_chinese(ch_num)}章', block):
                        continue
                    if re.match(rf'(?:Chapter|CHAPTER)\s+{ch_num}\b', block):
                        continue
                
                # Horizontal rule
                if block.startswith('---') or block.startswith('___'):
                    add_para(doc, '─' * 40, color=RGBColor(0xCC, 0xCC, 0xCC), size=8)
                    continue
                
                # Epigraph (starts with >)
                if block.startswith('>'):
                    epi_text = block.lstrip('>').strip()
                    # Handle multi-line epigraphs
                    lines = epi_text.split('\n')
                    for line in lines:
                        line = line.strip().lstrip('>').strip()
                        if line:
                            add_para(doc, line, italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x77, 0x77, 0x77))
                    doc.add_paragraph()
                    continue
                
                # Regular text - remove markdown formatting
                clean_block = block
                # Remove ** markers
                clean_block = re.sub(r'\*\*([^*]+)\*\*', r'\1', clean_block)
                # Remove __ markers
                clean_block = re.sub(r'__([^_]+)__', r'\1', clean_block)
                # Remove markdown headings
                clean_block = re.sub(r'^#{1,6}\s+', '', clean_block, flags=re.MULTILINE)
                # Remove code block markers
                clean_block = re.sub(r'```[\w]*\n?', '', clean_block)
                
                if clean_block.strip():
                    add_para(doc, clean_block.strip(), first_indent=(not first_para))
                    first_para = False
            
            added_count += 1
        else:
            # Fallback: use chapter synopsis
            add_para(doc, f'（{title_cn} - {title_en}）', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x99, 0x99, 0x99))
        
        # Add page break between chapters
        doc.add_page_break()
    
    # Save
    doc.save(OUTPUT)
    print(f"\nComplete! {added_count}/50 chapters added.")
    print(f"Output: {OUTPUT}")

def to_chinese(n):
    """Convert 1-100 to Chinese numeral string."""
    cn = ['零', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十',
          '十一', '十二', '十三', '十四', '十五', '十六', '十七', '十八', '十九', '二十',
          '二十一', '二十二', '二十三', '二十四', '二十五', '二十六', '二十七', '二十八', '二十九', '三十',
          '三十一', '三十二', '三十三', '三十四', '三十五', '三十六', '三十七', '三十八', '三十九', '四十',
          '四十一', '四十二', '四十三', '四十四', '四十五', '四十六', '四十七', '四十八', '四十九', '五十']
    if 1 <= n <= 50:
        return cn[n]
    # For 50+
    tens = n // 10
    ones = n % 10
    t_char = ['', '十', '二十', '三十', '四十', '五十', '六十'][tens]
    o_char = ['', '一', '二', '三', '四', '五', '六', '七', '八', '九'][ones]
    return t_char + o_char

if __name__ == '__main__':
    build_novel()
