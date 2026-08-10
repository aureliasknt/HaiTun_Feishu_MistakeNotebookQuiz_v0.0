#!/usr/bin/env python
"""Build the complete novel .docx from all chapter files."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
import re

OUTPUT = r"D:\haitun agent\psi-agent\generated\novel\星穹遗档_The_Archive_of_the_Starry_Dome.docx"

def add_styled_heading(doc, text, level=1):
    """Add a heading with proper styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_chapter(doc, title_cn, title_en, chapter_num, text_content):
    """Add a chapter to the document."""
    # Chapter heading
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"第{to_chinese_num(chapter_num)}章  {title_cn}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"Chapter {chapter_num}: {title_en}")
    run2.font.size = Pt(11)
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    # Add a page break before each chapter (except chapter 1)
    if chapter_num > 1:
        doc.add_page_break()
    
    # Parse and add the content
    # Process the text to handle paragraphs
    paragraphs = text_content.strip().split('\n\n')
    
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue
        
        # Check if it's a chapter heading (already handled above)
        if para_text.startswith('##') or para_text.startswith('# '):
            continue
            
        # Check if it's an epigraph (italic block)
        if para_text.startswith('>'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lines = para_text.strip('>').strip().split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    run = p.add_run(line + '\n')
                    run.font.italic = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            doc.add_paragraph()  # spacer
        elif para_text.startswith('---'):
            # Horizontal rule - add a spacer
            p = doc.add_paragraph()
            run = p.add_run('─' * 40)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
        else:
            # Regular paragraph
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.75)
            p.paragraph_format.line_spacing = Pt(22)
            run = p.add_run(para_text)
            run.font.size = Pt(11)
            run.font.name = 'Calibri'

def to_chinese_num(n):
    """Convert number to Chinese numerals."""
    cn = ['零', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
    if n <= 10:
        return cn[n]
    elif n < 20:
        return '十' + (cn[n - 10] if n > 10 else '')
    elif n < 100:
        tens = n // 10
        ones = n % 10
        return cn[tens] + '十' + (cn[ones] if ones > 0 else '')
    return str(n)

def setup_styles(doc):
    """Configure document styles for Chinese text."""
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = Pt(22)
    
    # Set east-asian font
    from docx.oxml.ns import qn
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = doc.oxml.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '微软雅黑')

def create_title_page(doc):
    """Create the title page."""
    for _ in range(6):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('星穹遗档')
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('The Archive of the Starry Dome')
    run2.font.size = Pt(18)
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x77)
    
    doc.add_paragraph()
    
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run('—— 一部跨越星辰的科幻史诗 ——')
    run3.font.size = Pt(12)
    run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    
    for _ in range(8):
        doc.add_paragraph()
    
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run('全书共50章 · 三部曲')
    run4.font.size = Pt(11)
    run4.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_page_break()

def create_toc_page(doc):
    """Create a table of contents page."""
    add_styled_heading(doc, '目 录', level=1)
    
    toc = [
        ('第一部：遗档之秘', 'The Secret of the Archive', '第1-18章'),
        ('', '第一章  最后一代档案员', '3'),
        ('', '第二章  痕迹', '15'),
        ('', '第三章  四百年', '28'),
        ('', '第四章  学徒', '40'),
        ('', '第五章  加密层', '52'),
        ('', '第六章  来自远方', '65'),
        ('', '第七章  陌生人', '78'),
        ('', '第八章  裂隙', '91'),
        ('', '第九章  鹤的沉默', '104'),
        ('', '第十章  抉择', '118'),
        ('', '第十一章  后门', '132'),
        ('', '第十二章  数据深渊', '145'),
        ('', '第十三章  静默者', '158'),
        ('', '第十四章  入侵', '172'),
        ('', '第十五章  两面作战', '185'),
        ('', '第十六章  征服者的欲望', '198'),
        ('', '第十七章  破碎', '212'),
        ('', '第十八章  逃亡协议', '225'),
        ('第二部：光速暗影', 'Shadows of Light-Speed', '第19-36章'),
        ('', '第十九章  流亡', '239'),
        ('', '第二十章  静默者的历史', '252'),
        ('', '第二十一章  三百年', '265'),
        ('', '第二十二章  鹤的告白', '278'),
        ('', '第二十三章  追踪', '291'),
        ('', '第二十四章  加密碎片', '304'),
        ('', '第二十五章  分歧', '317'),
        ('', '第二十六章  战火', '330'),
        ('', '第二十七章  逃出天枢七', '343'),
        ('', '第二十八章  星际旅途', '356'),
        ('', '第二十九章  荒废站', '369'),
        ('', '第三十章  死者的遗言', '382'),
        ('', '第三十一章  以太的本质', '395'),
        ('', '第三十二章  追逐者', '408'),
        ('', '第三十三章  影子战争', '421'),
        ('', '第三十四章  陷阱', '434'),
        ('', '第三十五章  双重夹击', '447'),
        ('', '第三十六章  绝境中的光芒', '460'),
        ('第三部：星辰抉择', 'Choice of the Stars', '第37-50章'),
        ('', '第三十七章  最后的信号', '475'),
        ('', '第三十八章  静默者之船', '488'),
        ('', '第三十九章  封印的历史', '501'),
        ('', '第四十章  代价', '514'),
        ('', '第四十一章  殊途', '527'),
        ('', '第四十二章  苏婉的决定', '540'),
        ('', '第四十三章  高孟德的最后一击', '553'),
        ('', '第四十四章  三方会谈', '566'),
        ('', '第四十五章  致命失误', '579'),
        ('', '第四十六章  牺牲', '592'),
        ('', '第四十七章  告别鹤', '605'),
        ('', '第四十八章  第三条路', '618'),
        ('', '第四十九章  散场', '631'),
        ('', '第五十章  星辰有光', '644'),
    ]
    
    for item in toc:
        p = doc.add_paragraph()
        if item[0]:
            run = p.add_run(item[0])
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            run2 = p.add_run(f'  {item[1]}')
            run2.font.size = Pt(10)
            run2.font.italic = True
            run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        else:
            run = p.add_run(f'    {item[1]}')
            run.font.size = Pt(10)
    
    doc.add_page_break()

def main():
    doc = Document()
    setup_styles(doc)
    
    create_title_page(doc)
    create_toc_page(doc)
    
    # Part 1 heading
    add_styled_heading(doc, '第一部：遗档之秘', level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('The Secret of the Archive')
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()
    
    # Now read each chapter file and add it
    chapters_dir = r"D:\haitun agent\psi-agent\generated\novel\chapters"
    
    chapter_files = [
        ('part1a_ch1_10.md', 1, 10),
        ('part1b_ch11_18.md', 11, 18),
        ('part2a_ch19_27.md', 19, 27),
        ('part2b_ch28_36.md', 28, 36),
        ('part3_ch37_50.md', 37, 50),
    ]
    
    # Chinese/English chapter title mapping
    chapter_titles = {
        1: ('最后一代档案员', 'The Last Archivist'),
        2: ('痕迹', 'Traces'),
        3: ('四百年', 'Four Hundred Years'),
        4: ('学徒', 'The Apprentice'),
        5: ('加密层', 'The Encryption Layer'),
        6: ('来自远方', 'From Afar'),
        7: ('陌生人', 'The Stranger'),
        8: ('裂隙', 'Cracks'),
        9: ('鹤的沉默', 'Hé\'s Silence'),
        10: ('抉择', 'The Choice'),
        11: ('后门', 'The Backdoor'),
        12: ('数据深渊', 'The Data Abyss'),
        13: ('静默者', 'The Silent Ones'),
        14: ('入侵', 'Invasion'),
        15: ('两面作战', 'Two Fronts'),
        16: ('征服者的欲望', 'The Conqueror\'s Desire'),
        17: ('破碎', 'Shattered'),
        18: ('逃亡协议', 'Escape Protocol'),
        19: ('流亡', 'Exile'),
        20: ('静默者的历史', 'History of the Silent Ones'),
        21: ('三百年', 'Three Hundred Years'),
        22: ('鹤的告白', 'Hé\'s Confession'),
        23: ('追踪', 'Pursuit'),
        24: ('加密碎片', 'Encrypted Fragments'),
        25: ('分歧', 'Division'),
        26: ('战火', 'Warfire'),
        27: ('逃出天枢七', 'Escape from Tianshu-7'),
        28: ('星际旅途', 'Voyage'),
        29: ('荒废站', 'The Derelict Station'),
        30: ('死者的遗言', 'The Dead\'s Last Words'),
        31: ('以太的本质', 'Nature of the Ether'),
        32: ('追逐者', 'The Pursuer'),
        33: ('影子战争', 'Shadow War'),
        34: ('陷阱', 'The Trap'),
        35: ('双重夹击', 'Between Two Fires'),
        36: ('绝境中的光芒', 'Light in Despair'),
        37: ('最后的信号', 'The Last Signal'),
        38: ('静默者之船', 'The Silent Ones\' Ship'),
        39: ('封印的历史', 'Sealed History'),
        40: ('代价', 'The Price'),
        41: ('殊途', 'Divergent Paths'),
        42: ('苏婉的决定', 'Su Wan\'s Decision'),
        43: ('高孟德的最后一击', 'Gao Mengde\'s Final Strike'),
        44: ('三方会谈', 'Three-Way Parley'),
        45: ('致命失误', 'Fatal Error'),
        46: ('牺牲', 'Sacrifice'),
        47: ('告别鹤', 'Farewell to Hé'),
        48: ('第三条路', 'The Third Path'),
        49: ('散场', 'Departure'),
        50: ('星辰有光', 'The Stars Still Shine'),
    }
    
    for file_name, start_ch, end_ch in chapter_files:
        file_path = os.path.join(chapters_dir, file_name)
        
        if os.path.exists(file_path):
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Try to split by chapters
            chapter_pattern = re.compile(r'(?:第[^章]*章[：:]?\s*第?\d+[^\n]*|Chapter\s+\d+)', re.IGNORECASE)
            
            # Add chapters from this file
            for ch_num in range(start_ch, end_ch + 1):
                title_cn, title_en = chapter_titles.get(ch_num, ('', ''))
                
                # Try to extract this chapter's content
                # Simple approach: look for chapter marker
                ch_marker = f"## 第{to_chinese_num(ch_num)}章"
                alt_marker = f"Chapter {ch_num}"
                
                ch_start = content.find(ch_marker)
                if ch_start < 0:
                    ch_start = content.find(alt_marker)
                
                next_ch_marker = None
                if ch_num < end_ch:
                    next_ch_marker_1 = f"## 第{to_chinese_num(ch_num+1)}章"
                    next_ch_marker_2 = f"## {to_chinese_num(ch_num+1)}"
                    next_ch_start = content.find(next_ch_marker_1, ch_start + 1)
                    if next_ch_start < 0:
                        next_ch_start = content.find(f"Chapter {ch_num+1}", ch_start + 1)
                else:
                    next_ch_start = len(content)
                
                if ch_start >= 0:
                    if next_ch_start and next_ch_start > ch_start:
                        chapter_text = content[ch_start:next_ch_start].strip()
                    else:
                        chapter_text = content[ch_start:].strip()
                    
                    add_chapter(doc, title_cn, title_en, ch_num, chapter_text)
        else:
            print(f"File not found: {file_path}")
    
    # Part 2 heading
    doc.add_page_break()
    add_styled_heading(doc, '第二部：光速暗影', level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Shadows of Light-Speed')
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()
    
    # Part 3 heading
    doc.add_page_break()
    add_styled_heading(doc, '第三部：星辰抉择', level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Choice of the Stars')
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()
    
    # Save
    doc.save(OUTPUT)
    print(f"Novel saved to: {OUTPUT}")

if __name__ == '__main__':
    main()
